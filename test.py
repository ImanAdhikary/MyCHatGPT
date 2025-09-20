import streamlit as st
from langgraph_backend import chatbot
from langchain_core.messages import HumanMessage
import datetime, json, os, tempfile
from pathlib import Path

# RAG imports
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from PyPDF2 import PdfReader
import docx
import openpyxl

# ---------- CONFIG ----------
CHAT_FILE = "chat_sessions.json"
AVATAR_USER = "🙂"
AVATAR_AI = "🤖"

# ---------- HELPERS ----------
def load_chats():
    if os.path.exists(CHAT_FILE):
        with open(CHAT_FILE, "r") as f:
            data = json.load(f)

        # --- MIGRATION STEP ---
        migrated = {}
        for tid, val in data.items():
            if isinstance(val, list):  # old format
                migrated[tid] = {
                    "title": "New Chat" if not val else " ".join(val[0]["content"].split()[:6]),
                    "messages": val
                }
            elif isinstance(val, dict):  # already new format
                migrated[tid] = val
        return migrated

    # default
    return {
        "thread-1": {
            "title": "New Chat",
            "messages": []
        }
    }


def save_chats():
    with open(CHAT_FILE, "w") as f:
        json.dump(st.session_state.chat_sessions, f)

def export_chat(thread_id):
    history = st.session_state.chat_sessions[thread_id]["messages"]
    if not history:
        return ""
    export_text = "\n".join(
        [f"[{msg['timestamp']}] {msg['role'].upper()}: {msg['content']}" for msg in history]
    )
    Path(f"{thread_id}.txt").write_text(export_text, encoding="utf-8")
    return f"{thread_id}.txt"

def process_uploaded_file(uploaded_file):
    """Convert uploaded file into FAISS retriever"""
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name

    raw_text = ""

    if uploaded_file.name.endswith(".pdf"):
        pdf = PdfReader(tmp_path)
        for page in pdf.pages:
            raw_text += page.extract_text() + "\n"

    elif uploaded_file.name.endswith(".docx"):
        doc = docx.Document(tmp_path)
        for para in doc.paragraphs:
            raw_text += para.text + "\n"

    elif uploaded_file.name.endswith(".xlsx"):
        wb = openpyxl.load_workbook(tmp_path)
        sheet = wb.active
        for row in sheet.iter_rows(values_only=True):
            raw_text += " ".join([str(cell) for cell in row if cell]) + "\n"

    elif uploaded_file.name.endswith(".txt"):
        with open(tmp_path, "r", encoding="utf-8") as f:
            raw_text = f.read()

    else:
        st.error("❌ Unsupported file type")
        return None

    # Split & embed
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = [Document(page_content=chunk) for chunk in splitter.split_text(raw_text)]
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    vectorstore = FAISS.from_documents(docs, embeddings)

    return vectorstore.as_retriever(search_kwargs={"k": 3})

def summarize_text(text, length=6):
    """Return a short title (5–6 words) from first user input"""
    words = text.split()
    return " ".join(words[:length]) if words else "New Chat"

# ---------- SESSION ----------
if "chat_sessions" not in st.session_state:
    st.session_state.chat_sessions = load_chats()
if "active_thread" not in st.session_state:
    st.session_state.active_thread = "thread-1"
if "retriever" not in st.session_state:
    st.session_state.retriever = None

# ---------- SIDEBAR ----------
st.sidebar.title("⚙️ Chat Controls")

# Resume chat
available_threads = list(st.session_state.chat_sessions.keys())
thread_labels = [
    f"{st.session_state.chat_sessions[tid]['title']}" for tid in available_threads
]
selected_index = available_threads.index(st.session_state.active_thread)
selected_label = st.sidebar.selectbox("Resume Chat:", thread_labels, index=selected_index)
selected_thread = available_threads[thread_labels.index(selected_label)]

if selected_thread != st.session_state.active_thread:
    st.session_state.active_thread = selected_thread

# New Chat
if st.sidebar.button("➕ New Chat"):
    new_thread_id = f"thread-{len(st.session_state.chat_sessions)+1}"
    st.session_state.chat_sessions[new_thread_id] = {"title": "New Chat", "messages": []}
    st.session_state.active_thread = new_thread_id
    save_chats()

# Export Chat
if st.sidebar.button("⬇️ Export Chat"):
    file_path = export_chat(st.session_state.active_thread)
    st.sidebar.success(f"Exported to {file_path}")

# Clear Chats
if st.sidebar.button("🗑️ Clear All Chats"):
    st.session_state.chat_sessions = {
        "thread-1": {"title": "New Chat", "messages": []}
    }
    st.session_state.active_thread = "thread-1"
    save_chats()
    st.sidebar.warning("All chats cleared!")

# Observability
st.sidebar.subheader("🔍 Observability")
if st.sidebar.checkbox("Show Logs"):
    for tid, data in st.session_state.chat_sessions.items():
        st.sidebar.markdown(f"**{data['title']}** ({len(data['messages'])} msgs)")
        for msg in data["messages"][-5:]:  # last 5 only
            st.sidebar.caption(f"[{msg['timestamp']}] {msg['role']}: {msg['content'][:50]}...")

# File Upload for RAG
st.sidebar.subheader("📂 Upload Knowledge")
uploaded_file = st.sidebar.file_uploader("Upload PDF/Doc/Excel/TXT", type=["pdf", "docx", "xlsx", "txt"])
if uploaded_file:
    st.session_state.retriever = process_uploaded_file(uploaded_file)
    if st.session_state.retriever:
        st.sidebar.success(f"✅ {uploaded_file.name} loaded into RAG")
    else:
        st.sidebar.error("⚠️ Could not process file")

# ---------- MAIN CHAT ----------
st.title("💬 Next-Level AI Chat with RAG")

for message in st.session_state.chat_sessions[st.session_state.active_thread]["messages"]:
    avatar = AVATAR_USER if message["role"] == "user" else AVATAR_AI
    with st.chat_message(message["role"], avatar=avatar):
        st.markdown(message["content"])

# ---------- USER INPUT ----------
user_input = st.chat_input("Type here...")

if user_input:
    # Add user message
    user_msg = {
        "role": "user",
        "content": user_input,
        "timestamp": datetime.datetime.now().strftime("%H:%M:%S")
    }
    st.session_state.chat_sessions[st.session_state.active_thread]["messages"].append(user_msg)

    # If first message in this thread → set title
    if st.session_state.chat_sessions[st.session_state.active_thread]["title"] == "New Chat":
        st.session_state.chat_sessions[st.session_state.active_thread]["title"] = summarize_text(user_input)

    save_chats()

    with st.chat_message("user", avatar=AVATAR_USER):
        st.markdown(user_input)

    # RAG context
    context = ""
    if st.session_state.retriever:
        docs = st.session_state.retriever.get_relevant_documents(user_input)
        if docs:
            context = "\n\n".join([d.page_content for d in docs])

    full_prompt = f"Context:\n{context}\n\nQuestion: {user_input}"

    # AI response
    with st.chat_message("assistant", avatar=AVATAR_AI):
        ai_message = st.write_stream(
            message_chunk.content for message_chunk, metadata in chatbot.stream(
                {"messages": [HumanMessage(content=full_prompt)]},
                config={"configurable": {"thread_id": st.session_state.active_thread}},
                stream_mode="messages"
            )
        )


    ai_msg = {
        "role": "assistant",
        "content": ai_message,
        "timestamp": datetime.datetime.now().strftime("%H:%M:%S")
    }
    st.session_state.chat_sessions[st.session_state.active_thread]["messages"].append(ai_msg)
    save_chats()
